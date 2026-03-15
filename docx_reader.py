"""Извлечение текста из .docx для передачи в модель."""
from pathlib import Path

from docx import Document


def extract_text(docx_path: Path) -> str:
    """
    Извлекает весь текст из .docx: параграфы и таблицы.
    """
    doc = Document(docx_path)
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n\n".join(parts) if parts else ""

